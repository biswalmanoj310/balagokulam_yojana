import docx
from read_create_file_folder import read_system_value, create_main_folder, create_genai_summary_file
from topic_summary_gen import generate_topic_summary


# system_info_file_path = 'system_new_info.txt'
system_info_file_path = 'system_information.txt'
system_info = read_system_value(system_info_file_path)

with open('All_Topics.txt', 'r') as file:
    topic_names = file.readlines()

"""
topic_names = ["Sages and Saints from Satya Yuga", "Sages and Saints from Tretaya Yuga",
              ]
"""

ollama_model_name = "llama3:latest"
# ollama_model_name = "wizardlm2:7b"
# ollama_model_name = "gemma2:9b"

# ollama_model_names = ["llama3:latest", "wizardlm2:7b"]

# for ollama_model_name in ollama_model_names:
for topic_name_full in topic_names:
    topic_name = topic_name_full.strip()
    output_text = generate_topic_summary(system_info, topic_name, ollama_model_name)

    document = docx.Document()
    document.add_heading(topic_name, 0)
    document.add_paragraph(output_text)

    create_main_folder(topic_name)

    create_genai_summary_file(topic_name, document)

